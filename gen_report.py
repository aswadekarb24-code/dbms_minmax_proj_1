from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_report():
    doc = Document()

    # --- STYLE ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- TITLE PAGE ---
    title = doc.add_heading('Project Report: TPQA Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Academic Consultancy & Third-Party Quality Audit Management Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('DBMS Mini Project\n').bold = True
    info.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
    info.add_run('Institution: VJTI, Mumbai\n')

    doc.add_page_break()

    # --- PROBLEM STATEMENT ---
    doc.add_heading('1. Problem Statement', level=1)
    doc.add_paragraph(
        "Academic consultancy and Third-Party Quality Audit (TPQA) projects involve complex workflows including "
        "multiple stakeholders like external organizations, institutional directors, heads of departments, and project coordinators. "
        "Historically, these processes were handled through manual paperwork, leading to several issues:"
    )
    doc.add_paragraph("Lack of real-time tracking of project status.", style='List Bullet')
    doc.add_paragraph("Inefficient revenue distribution between the institute and departments.", style='List Bullet')
    doc.add_paragraph("Difficulty in generating consistent proforma and tax invoices.", style='List Bullet')
    doc.add_paragraph("Potential for data inconsistency and human error in financial calculations.", style='List Bullet')
    doc.add_paragraph("Security concerns related to sensitive financial and institutional data.", style='List Bullet')

    # --- TOOLS & TECHNOLOGIES ---
    doc.add_heading('2. Tools and Technologies', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'

    technologies = [
        ('Frontend', 'Next.js 16, TailwindCSS 4, Axios'),
        ('Backend', 'FastAPI, SQLAlchemy, Pydantic'),
        ('Database', 'PostgreSQL (Supabase)'),
        ('Authentication', 'JWT (JSON Web Tokens), Bcrypt Hashing'),
        ('Testing', 'Pytest, HTTPX'),
        ('Deployment', 'Vercel (Frontend), Render/Supabase (Backend/DB)')
    ]

    for component, tech in technologies:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech

    # --- DATABASE SCHEMA ---
    doc.add_heading('3. Database Schema', level=1)
    doc.add_paragraph(
        "The system utilizes a relational PostgreSQL database with 11 primary tables and custom enumerations to manage "
        "the project lifecycle. Key components include:"
    )
    
    # Helper to add formatted bullets with optional indentations
    def add_bullet(text, bold_prefix="", indent=0):
        p = doc.add_paragraph(style='List Bullet')
        if indent > 0:
            p.paragraph_format.left_indent = Inches(0.5 * indent)
        if bold_prefix:
            p.add_run(bold_prefix).bold = True
        p.add_run(text)

    add_bullet(" Manages faculty members (Director, HOD, Coordinator) and their permissions.", "Employees & Roles:")
    add_bullet(" Enforces role-based access control with predefined constraints checking for DIRECTOR, HOD, PROJECT_COORDINATOR, and SUPPORT_STAFF.", "roles:", 1)
    add_bullet(" Centralizes user profiles by linking a user to a specific department_id and role_id, while maintaining secure fields like auth_hash, email, and tracking their professional development balance (pdf_balance).", "employees:", 1)

    add_bullet(" Stores information about external organizations requesting consultancy.", "Clients & Agencies:")
    add_bullet(" Records primary organizational profiles containing critical billing and contact data such as gstin_uin, state_code, and contact_email.", "clients:", 1)
    add_bullet(" Maintains records of sub-contracted third parties along with their services_scope.", "external_agencies:", 1)

    add_bullet(" Maps institutional divisions to specific users, linking each department_name to an hod_employee_id.", "Departments:")

    add_bullet(" The core entity tracking the 9-step workflow state. It acts as the central hub connecting client_id, department_id, and coordinator_id. It stores comprehensive logistics, including cost_of_work, est_person_days, and actively updates the current_status enumeration.", "Projects:")

    add_bullet(" Tables for Budget Estimations, Invoices, and Payment Receipts.", "Budget & Financials:")
    add_bullet(" Captures highly granular financial proposals tied to a project, isolating components like faculty_fees, cpts_charges, operational_exp, net_project_cost, and tracking explicit director_approval flags.", "budget_estimations:", 1)
    add_bullet(" Archives formalized billing data such as hsn_sac_code, taxable_value, and tax_amount, supporting distinct invoice_type definitions.", "invoices:", 1)
    add_bullet(" Logs specific payment transactions against invoices, verifying tds_deducted, transaction modes, and preventing duplicate entries via a unique bank_trans_ref.", "receipts:", 1)

    add_bullet(" Master-detail tables for the 70/30 revenue split between the Institute and Faculty.", "Distribution:")
    add_bullet(" Aggregates the broader project split based on a specific receipt, enforcing constraints for staff_pool_70, inst_pool_30, and recording final approval_status.", "distribution_master:", 1)
    add_bullet(" Executes granular fund assignments, mapping specific allocated_amt and percentage_rule distributions to individual employees or external agencies depending on the payee_type.", "distribution_line_items:", 1)

    # --- SOLUTION IDEA ---
    doc.add_heading('4. Solution Idea', level=1)
    doc.add_paragraph(
        "The solution is built around a linear 9-step pipeline that enforces the institutional workflow rules. "
        "Each step requires specific actions from either the Client or the Faculty, ensuring accountability."
    )
    
    steps = [
        "Step 1: Resource request by External Organization.",
        "Step 2: Faculty response (Budget Proposal & Coordinator Assignment).",
        "Step 3: Agency acceptance of the fee structure.",
        "Step 4: Institutional Approval by the Director.",
        "Step 5: Generation of Proforma Invoice.",
        "Step 6: Recording of Tax Invoice and Receipt of Payment.",
        "Step 7: Submission of Project Completion Reports.",
        "Step 8: Automated Revenue Distribution (70% Institute / 30% Department).",
        "Step 9: Project Closure."
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # --- CHALLENGES FACED ---
    doc.add_heading('5. Challenges Faced', level=1)
    challenges = [
        ("State Management: ", "Implementing a robust 9-step workflow where each step depends on the data from previous steps while allowing for potential errors."),
        ("Financial Accuracy: ", "Ensuring precise calculations for GST, TDS, and institutional revenue shares across various financial documents."),
        ("Security & RBAC: ", "Designing a Role-Based Access Control system that correctly partitions data (e.g., HODs only see their department's projects)."),
        ("Error Recovery: ", "Developing a 'Step Reversal' mechanism to allow coordinators to correct mistakes in previous steps without corrupting the database state."),
        ("UI/UX for Complex Forms: ", "Displaying large amounts of financial data in a user-friendly way using modern Next.js components.")
    ]
    for title_text, desc in challenges:
        add_bullet(desc, title_text)

    # --- FUTURE IMPROVEMENTS ---
    doc.add_heading('6. Future Improvements', level=1)
    improvements = [
        ("Automated Notifications: ", "Integrating SMTP or SMS services to notify stakeholders when their action is required."),
        ("Mobile App: ", "Developing a React Native application for directors to approve budgets on the go."),
        ("ERP Integration: ", "Connecting with the college's existing SAP/ERP system for seamless payroll and accounting."),
        ("AI Analytics: ", "Using machine learning to predict budget trends and resource utilization based on historical data."),
        ("Digital Certificates: ", "Generating blockchain-verified completion certificates for consultancy projects.")
    ]
    for title_text, desc in improvements:
        add_bullet(desc, title_text)

    # --- CONCLUSION ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "The TPQA Management System successfully digitizes the complex consultancy workflow of an engineering institution. "
        "By providing transparency, automated calculations, and secure access, the system reduces administrative overhead "
        "and ensures financial integrity. The modern tech stack ensures scalability and a premium user experience."
    )

    doc.add_page_break()

    # --- DELIVERABLES & RESOURCES ---
    doc.add_heading('Project Deliverables & Resources', level=1)
    doc.add_paragraph("[Drive Link: Demo Video] [ __________________________________________________ ]").bold = True
    doc.add_paragraph()
    doc.add_paragraph("[Drive Link: ER Diagram SQL & Image] [ __________________________________________________ ]").bold = True
    doc.add_paragraph()
    doc.add_paragraph("[Drive Link: Workflow Diagram] [ __________________________________________________ ]").bold = True
    
    doc.add_page_break()

    # --- SCRIPTS ---
    doc.add_heading('Scripts', level=1)
    
    doc.add_heading('ER Diagram (Mermaid Script)', level=2)
    mermaid_er = (
        "erDiagram\n"
        "    clients ||--o{ projects : \"requests\"\n"
        "    departments ||--o{ employees : \"has\"\n"
        "    departments ||--o{ projects : \"owns\"\n"
        "    roles ||--o{ employees : \"defines\"\n"
        "    employees ||--o{ projects : \"coordinates\"\n"
        "    employees ||--o{ departments : \"heads (HOD)\"\n"
        "    employees ||--o{ distribution_line_items : \"receives\"\n"
        "    projects ||--o| budget_estimations : \"has\"\n"
        "    projects ||--o{ invoices : \"generates\"\n"
        "    projects ||--o{ distribution_master : \"requires\"\n"
        "    external_agencies ||--o{ budget_estimations : \"listed in\"\n"
        "    external_agencies ||--o{ distribution_line_items : \"receives\"\n"
        "    invoices ||--o{ receipts : \"paid via\"\n"
        "    receipts ||--o{ distribution_master : \"triggers\"\n"
        "    distribution_master ||--o{ distribution_line_items : \"contains\"\n\n"
        "    clients {\n"
        "        integer client_id PK\n"
        "        string organization_name\n"
        "        string contact_email UK\n"
        "        string gstin_uin\n"
        "    }\n"
        "    departments {\n"
        "        integer department_id PK\n"
        "        string department_name UK\n"
        "        integer hod_employee_id FK\n"
        "    }\n"
        "    roles {\n"
        "        integer role_id PK\n"
        "        string role_name UK\n"
        "    }\n"
        "    employees {\n"
        "        integer employee_id PK\n"
        "        integer department_id FK\n"
        "        integer role_id FK\n"
        "        string full_name\n"
        "        string email UK\n"
        "        numeric pdf_balance\n"
        "    }\n"
        "    external_agencies {\n"
        "        integer ext_agency_id PK\n"
        "        string agency_name\n"
        "        text services_scope\n"
        "    }\n"
        "    projects {\n"
        "        integer project_id PK\n"
        "        string project_number UK\n"
        "        integer client_id FK\n"
        "        integer department_id FK\n"
        "        integer coordinator_id FK\n"
        "        string project_title\n"
        "        string current_status\n"
        "        numeric cost_of_work\n"
        "    }\n"
        "    budget_estimations {\n"
        "        integer budget_id PK\n"
        "        integer project_id FK\n"
        "        integer ext_agency_id FK\n"
        "        numeric total_project_cost\n"
        "        boolean director_approval\n"
        "    }\n"
        "    invoices {\n"
        "        integer invoice_id PK\n"
        "        integer project_id FK\n"
        "        string invoice_number UK\n"
        "        date invoice_date\n"
        "        numeric total_amount\n"
        "    }\n"
        "    receipts {\n"
        "        integer receipt_id PK\n"
        "        integer invoice_id FK\n"
        "        string voucher_number UK\n"
        "        date receipt_date\n"
        "        numeric total_received\n"
        "        string bank_trans_ref UK\n"
        "    }\n"
        "    distribution_master {\n"
        "        integer dist_master_id PK\n"
        "        integer project_id FK\n"
        "        integer receipt_id FK\n"
        "        numeric total_dist_amt\n"
        "        numeric staff_pool_70\n"
        "        numeric inst_pool_30\n"
        "        boolean approval_status\n"
        "    }\n"
        "    distribution_line_items {\n"
        "        integer line_item_id PK\n"
        "        integer dist_master_id FK\n"
        "        integer employee_id FK\n"
        "        integer ext_agency_id FK\n"
        "        numeric percentage_rule\n"
        "        numeric allocated_amt\n"
        "    }"
    )
    p_er = doc.add_paragraph(mermaid_er)
    for run in p_er.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('Workflow Diagram (Mermaid Script)', level=2)
    mermaid_wf = (
        "flowchart TD\n"
        "    Start([Start Project]) --> Step1[Step 1: Resource request by External Organization]\n"
        "    Step1 --> Step2[Step 2: Faculty response - Budget & Coordinator]\n"
        "    Step2 --> Step3[Step 3: Agency acceptance of fee structure]\n"
        "    Step3 --> Step4[Step 4: Institutional Approval by Director]\n"
        "    Step4 --> Step5[Step 5: Generation of Proforma Invoice]\n"
        "    Step5 --> Step6[Step 6: Record Tax Invoice & Payment Receipt]\n"
        "    Step6 --> Step7[Step 7: Submission of Completion Reports]\n"
        "    Step7 --> Step8[Step 8: Automated Revenue Distribution]\n"
        "    Step8 --> Step9[Step 9: Project Closure]\n"
        "    Step9 --> End([End Project])\n\n"
        "    classDef step fill:#f9f9f9,stroke:#333,stroke-width:2px;\n"
        "    class Step1,Step2,Step3,Step4,Step5,Step6,Step7,Step8,Step9 step;"
    )
    p_wf = doc.add_paragraph(mermaid_wf)
    for run in p_wf.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # --- FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.text = "Generated by TPQA Management System - Internal Documentation"
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('Project_Report_TPQA.docx')
    print("Report generated successfully: Project_Report_TPQA.docx")

if __name__ == "__main__":
    create_report()